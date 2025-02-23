import streamlit as st
import openai
import docx
from docx import Document
import tempfile
import shutil
import difflib
from datetime import datetime

# -------------------
# CONFIGURAÇÃO OPENAI
# -------------------
API_KEY = "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"
openai.api_key = API_KEY

# -------------------
# FUNÇÃO PARA REVISAR TEXTO COM O GPT (USANDO A NOVA API OPENAI)
# -------------------
def review_text_with_gpt(text):
    prompt = (
        "You are an expert in academic writing. Improve the following text for grammar, clarity, and academic style without changing its meaning. "
        "Do not add comments or explanations—just return the improved text.\n\n"
        f"{text}"
    )

    client = openai.Client(api_key=API_KEY)

    response = client.chat.completions.create(
        model="gpt-4",
        messages=[
            {"role": "system", "content": "You are an expert academic editor."},
            {"role": "user", "content": prompt}
        ],
        temperature=0.1,
        max_tokens=1500,
    )

    return response.choices[0].message.content.strip()

# -------------------
# FUNÇÃO PARA ADICIONAR ALTERAÇÕES (TRACK CHANGES SIMULADO)
# -------------------
def apply_track_changes(original_text, revised_text):
    diff = difflib.ndiff(original_text.split(), revised_text.split())
    tracked_text = ""

    for word in diff:
        if word.startswith('- '):
            tracked_text += f"[DELETED: {word[2:]}] "
        elif word.startswith('+ '):
            tracked_text += f"[ADDED: {word[2:]}] "
        elif word.startswith('  '):
            tracked_text += f"{word[2:]} "

    return tracked_text.strip()

# -------------------
# FUNÇÃO PARA PROCESSAR O DOCUMENTO WORD
# -------------------
def process_word_document(uploaded_file):
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as temp_input:
        shutil.copyfileobj(uploaded_file, temp_input)

    doc = Document(temp_input.name)

    for para in doc.paragraphs:
        original_text = para.text.strip()
        if original_text:
            revised_text = review_text_with_gpt(original_text)
            tracked_text = apply_track_changes(original_text, revised_text)

            # Substitui o texto original pelo texto com alterações
            para.clear()
            para.add_run(tracked_text)

    # Salva o documento revisado
    output_path = tempfile.NamedTemporaryFile(delete=False, suffix="_tracked.docx").name
    doc.save(output_path)

    return output_path

# -------------------
# INTERFACE STREAMLIT
# -------------------
st.title("📄 MalquePub Editing Services - Track Changes")

st.write("""
Upload a Word document (.docx) and receive a version with grammar and academic style improvements. 
Corrections will be provided with tracked changes (additions and deletions highlighted). 
""")

uploaded_file = st.file_uploader("Upload your Word document:", type=["docx"])

if uploaded_file and st.button("🚀 Process Document"):
    with st.spinner("Processing... Please wait."):
        try:
            output_path = process_word_document(uploaded_file)
            st.success("✅ Document processed successfully!")

            with open(output_path, "rb") as file:
                st.download_button(
                    label="⬇️ Download Revised Document",
                    data=file,
                    file_name=f"revised_with_track_changes_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
        except Exception as e:
            st.error(f"❌ An error occurred: {str(e)}")
